from yt_dlp import YoutubeDL
from youtube_transcript_api import YouTubeTranscriptApi, NoTranscriptFound, TranscriptsDisabled
from docx import Document
import re
import os


class YouTubeTranscriptDownloader:
    def __init__(self, output_dir='transcripts', ffmpeg_path=None, language='pt-BR'):
        # Directory to save transcript files
        self.output_dir = output_dir
        self.ffmpeg_path = ffmpeg_path
        self.language = language  # Preferred language for the transcript or subtitle
        self.fallback_language = None  # Will be set dynamically

        # yt-dlp options
        self.ydl_opts = {
            'quiet': True,
            'skip_download': True,
            'writesubtitles': True,
            'subtitleslangs': [self.language],  # Preferred language
            'subtitlesformat': 'ttml',
            'outtmpl': f"{self.output_dir}/%(title)s.%(ext)s",
            'ffmpeg_location': self.ffmpeg_path,
        }

        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)

    @staticmethod
    def sanitize_filename(filename):
        # Replace invalid filename characters with underscores
        return re.sub(r'[\\/*?:"<>|]', '_', filename)

    def save_to_docx(self, content, title):
        """Save content to a .docx file."""
        sanitized_title = self.sanitize_filename(title)
        transcript_path = os.path.join(self.output_dir, f"{sanitized_title}.docx")
        doc = Document()
        doc.add_heading(title, 0)
        for line in content:
            doc.add_paragraph(line)
        doc.save(transcript_path)
        print(f"Transcript saved: {transcript_path}")

    def detect_original_language(self, video_url):
        """Detect the video's original language."""
        ydl_opts = {
            'quiet': True,
            'skip_download': True,
        }

        with YoutubeDL(ydl_opts) as ydl:
            video_info = ydl.extract_info(video_url, download=False)
            self.fallback_language = video_info.get('original_language', 'en')
            print(f"Detected original language: {self.fallback_language}")
            # Update ydl_opts with the fallback language
            self.ydl_opts['subtitleslangs'].append(self.fallback_language)

    def fetch_video_transcripts_or_captions(self, video_url):
        """Fetch and save the transcript or captions for a single video."""
        self.detect_original_language(video_url)  # Detect original language dynamically

        with YoutubeDL(self.ydl_opts) as ydl:
            video_info = ydl.extract_info(video_url, download=False)
            video_title = video_info.get('title', 'Unknown Title')
            video_id = video_info.get('id')

        try:
            # Attempt to get the transcript in the preferred language
            transcript = YouTubeTranscriptApi.get_transcript(video_id, languages=[self.language])
            content = [entry['text'] for entry in transcript]
            self.save_to_docx(content, video_title)
        except NoTranscriptFound:
            print(f"No transcript found for {video_title} in {self.language}. Trying fallback language...")
            # Fallback to the detected original language
            try:
                transcript = YouTubeTranscriptApi.get_transcript(video_id, languages=[self.fallback_language])
                content = [entry['text'] for entry in transcript]
                self.save_to_docx(content, video_title)
            except NoTranscriptFound:
                print(f"No transcript found for {video_title} in fallback language {self.fallback_language}.")
                self.fetch_subtitles(video_url, video_title)
        except TranscriptsDisabled:
            print(f"Transcripts are disabled for {video_title}. Trying captions...")
            self.fetch_subtitles(video_url, video_title)

    def fetch_subtitles(self, video_url, video_title):
        """Fetch and save subtitles (captions) using yt-dlp."""
        try:
            with YoutubeDL(self.ydl_opts) as ydl:
                info = ydl.extract_info(video_url, download=False)
                subtitle_path = os.path.join(self.output_dir, f"{self.sanitize_filename(video_title)}.ttml")

                if os.path.exists(subtitle_path):
                    print(f"Captions saved: {subtitle_path}")
                    # Convert subtitle file to text and save as .docx
                    with open(subtitle_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                    self.save_to_docx([content], video_title)
                    os.remove(subtitle_path)  # Clean up subtitle file
                else:
                    print(f"No captions available for {video_title}.")
        except Exception as e:
            print(f"Failed to fetch captions for {video_title}: {e}")

    def fetch_playlist_transcripts(self, playlist_url):
        """Fetch and save transcripts for all videos in a playlist."""
        playlist_opts = {
            'quiet': True,
            'extract_flat': True,
            'skip_download': True,
        }

        with YoutubeDL(playlist_opts) as ydl:
            playlist_info = ydl.extract_info(playlist_url, download=False)
            entries = playlist_info.get('entries', [])
            print(f"Found {len(entries)} videos in playlist.")

            for video in entries:
                video_url = video.get('url')
                self.fetch_video_transcripts_or_captions(video_url)

    def download_transcripts(self, url):
        """Determine if the URL is a playlist or video and fetch transcripts."""
        with YoutubeDL(self.ydl_opts) as ydl:
            info = ydl.extract_info(url, download=False)
            if 'entries' in info:  # Playlist detected
                self.fetch_playlist_transcripts(url)
            else:  # Single video detected
                self.fetch_video_transcripts_or_captions(url)

# Example usage:
# Initialize the downloader with a language parameter
downloader = YouTubeTranscriptDownloader(
    output_dir='A:/Gravações/transcript',
    ffmpeg_path="E:/PycharmProjects/transcriptor/transcriptor/ffmpeg/bin/ffmpeg.exe",
    language="pt"  # For a list use this link https://github.com/yt-dlp/yt-dlp/issues/2205
)
# Pass a single video URL or playlist URL
url = "https://www.youtube.com/watch?v=Pny70rNPJLk"
downloader.download_transcripts(url)
